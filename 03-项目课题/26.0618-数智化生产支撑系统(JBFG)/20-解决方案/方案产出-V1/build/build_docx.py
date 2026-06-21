"""
JBFG Proposal DOCX Builder — One-click build from MD to final deliverable.

Stages:
  1. Preprocess MD (strip YAML, convert WikiLinks, inject image refs)
  2. Pandoc convert with reference template (--reference-doc)
  3. Post-process DOCX (cover page, headers/footers, images, table styling)
  4. Quality verification

Usage:
  python build_docx.py           # Full build
  python build_docx.py --check   # Verify existing DOCX only
"""

import os
import re
import sys
import subprocess
import shutil
from pathlib import Path

# ── Paths ──
BASE_DIR = Path(__file__).resolve().parent.parent  # 方案产出-V1/
BUILD_DIR = BASE_DIR / 'build'
IMAGES_DIR = BASE_DIR / 'images'
MD_SOURCE = BASE_DIR / '冀北风光储数智化生产支撑与数字孪生解决方案-V1.md'
MD_PREPROCESSED = BUILD_DIR / 'source_preprocessed.md'
REFERENCE_DOCX = BUILD_DIR / '参考模板.docx'
OUTPUT_DOCX = BASE_DIR / '冀北风光储数智化生产支撑与数字孪生解决方案-V2.docx'


# ══════════════════════════════════════════════
#  Stage 1: MD Preprocessing
# ══════════════════════════════════════════════

# Figure mapping: ASCII block identifier → image file
FIGURE_MAP = {
    '方案核心逻辑': ('images/fig1-1.png', '图 1-1：方案核心逻辑——三大矛盾通过 USDF 三通道分流统一解决'),
    'SCQA 分析框架': ('images/fig2-1.png', '图 2-1：SCQA 分析框架——从现状到方案的结构化推导'),
    '五大业务子系统 —— 建·控·评·检·管 闭环全景图': ('images/fig4-0.png', '图 4-0：五大业务子系统闭环全景——建→控→评→检→管，数字孪生驱动'),
    '四层 SLA 延迟预算与数据通路': ('images/fig5-1.png', '图 5-1：四层 SLA 延迟预算——从毫秒级暂态控制到秒级历史查询，每层独立通道、独立验收'),
    '项目实施甘特图': ('images/fig6-1.png', '图 6-1：项目实施甘特图——底座先行(1-2月) → 业务渐进(2-3月) → 联调上线(3-4月)'),
    '风险热力图': ('images/fig7-1.png', '图 7-1：风险热力图——18 项风险按技术/业务/实施三维度 × 三色等级分布'),
    '报价构成': ('images/fig8-1.png', '图 8-1：报价构成——220 万含税分配'),
}


def strip_yaml_frontmatter(content):
    """Remove YAML frontmatter (--- ... ---) from Markdown."""
    # Only match at the very beginning
    if content.startswith('---'):
        end = content.find('---', 3)
        if end != -1:
            return content[end + 3:].lstrip('\n')
    return content


def convert_wikilinks(content):
    """Convert Obsidian WikiLinks to plain text.
    [[link|display]] → display
    [[link]] → link
    [[link#anchor]] → link
    """
    def replace_link(m):
        inner = m.group(1)
        if '|' in inner:
            return inner.split('|', 1)[1]
        elif '#' in inner:
            return inner.split('#', 1)[0]
        return inner
    return re.sub(r'\[\[([^\]]+)\]\]', replace_link, content)


def replace_ascii_diagrams(content):
    """Replace ASCII art code blocks with image references for 7 key figures.
    Other code blocks (architecture diagrams, sequence diagrams) are preserved as code.
    """
    # Split content into segments: text + code blocks
    # We look for ``` ... ``` patterns
    parts = []
    i = 0
    lines = content.split('\n')

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```') and not line.strip().startswith('```mermaid'):
            # Start of code block
            block_start = i
            block_lines = [line]
            i += 1
            # Collect until closing ```
            while i < len(lines) and not lines[i].strip().startswith('```'):
                block_lines.append(lines[i])
                i += 1
            if i < len(lines):
                block_lines.append(lines[i])  # closing ```
                i += 1

            block_text = '\n'.join(block_lines)
            block_content = '\n'.join(block_lines[1:-1])  # content between ```

            # Check if this block matches any figure
            replaced = False
            for keyword, (img_path, caption) in FIGURE_MAP.items():
                if keyword in block_content:
                    # Replace ASCII block with image reference
                    parts.append(f'![{caption}]({img_path})')
                    parts.append('')
                    parts.append(f'> {caption}')
                    replaced = True
                    break

            if not replaced:
                parts.append(block_text)
        else:
            parts.append(line)
            i += 1

    return '\n'.join(parts)


def preprocess_md():
    """Preprocess the MD source file for Pandoc conversion."""
    print('[1/4] Preprocessing MD source...')

    with open(MD_SOURCE, 'r', encoding='utf-8') as f:
        content = f.read()

    # 1. Strip YAML frontmatter
    content = strip_yaml_frontmatter(content)
    print('  - YAML frontmatter stripped')

    # 2. Convert WikiLinks
    content = convert_wikilinks(content)
    print('  - WikiLinks converted')

    # 3. Replace ASCII diagrams with images (using absolute paths for Pandoc)
    content = replace_ascii_diagrams(content)
    # Convert relative image paths to absolute so Pandoc can find them
    images_abs = IMAGES_DIR.resolve()
    content = re.sub(
        r'!\[([^\]]*)\]\(images/([^)]+)\)',
        lambda m: f'![{m.group(1)}]({images_abs / m.group(2)})',
        content
    )
    print('  - ASCII diagrams → image references (absolute paths)')

    # 4. Remove Obsidian comments (%% ... %%)
    content = re.sub(r'%%.*?%%', '', content, flags=re.DOTALL)
    print('  - Comments removed')

    # Write preprocessed file
    with open(MD_PREPROCESSED, 'w', encoding='utf-8') as f:
        f.write(content)

    print(f'  → Saved: {MD_PREPROCESSED}')
    print(f'  → Lines: {len(content.splitlines())}')
    return MD_PREPROCESSED


# ══════════════════════════════════════════════
#  Stage 2: Pandoc Conversion
# ══════════════════════════════════════════════

def run_pandoc():
    """Convert preprocessed MD to DOCX using reference template."""
    print('\n[2/4] Running Pandoc conversion...')

    cmd = [
        'pandoc',
        str(MD_PREPROCESSED),
        '--from=gfm',
        '--to=docx',
        f'--reference-doc={REFERENCE_DOCX}',
        '--toc',
        '--toc-depth=3',
        '--metadata', 'title=冀北风光储数智化生产支撑与数字孪生解决方案',
        '--metadata', 'author=ZSSolutions',
        '--metadata', 'subject=新能源场站数智化生产支撑系统技术方案书',
        '--metadata', 'lang=zh-CN',
        '--metadata', 'keywords=冀北风光储,数智化,数字孪生,USDF,构网型,GOOSE',
        '-o', str(OUTPUT_DOCX),
    ]

    # Check if pandoc is available
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8',
                                errors='replace', timeout=60)
    except FileNotFoundError:
        print('  ERROR: Pandoc not found. Please install pandoc first.')
        return False
    except subprocess.TimeoutExpired:
        print('  ERROR: Pandoc timed out.')
        return False

    if result.returncode != 0:
        print(f'  ERROR: Pandoc failed:\n{result.stderr}')
        return False

    print(f'  → Saved: {OUTPUT_DOCX}')
    print(f'  → Size: {os.path.getsize(OUTPUT_DOCX):,} bytes')
    return True


# ══════════════════════════════════════════════
#  Stage 3: DOCX Post-processing
# ══════════════════════════════════════════════

def postprocess_docx():
    """Post-process the Pandoc-generated DOCX:
    - Add cover page
    - Configure headers/footers
    - Embed images
    - Style tables
    """
    print('\n[3/4] Post-processing DOCX...')

    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        from docx.oxml.shared import OxmlElement
    except ImportError:
        print('  ERROR: python-docx not installed')
        return False

    # Re-use color palette
    C_PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
    C_SECONDARY = RGBColor(0x2B, 0x5C, 0x8A)
    C_BODY = RGBColor(0x33, 0x33, 0x33)
    C_GRAY = RGBColor(0x66, 0x66, 0x66)
    C_TABLE_HDR = RGBColor(0xD6, 0xE4, 0xF0)
    BORDER_COLOR = 'CCCCCC'

    FONT_CN = '宋体'
    FONT_H = '黑体'
    FONT_EN = 'Times New Roman'

    doc = Document(str(OUTPUT_DOCX))

    # ── 3a. Add Cover Page ──
    print('  - Adding cover page...')
    _add_cover_page(doc, C_PRIMARY, C_SECONDARY, C_BODY, C_GRAY, FONT_CN, FONT_H)

    # ── 3b. Configure Headers / Footers ──
    print('  - Configuring headers/footers...')
    _configure_headers_footers(doc, C_GRAY, FONT_CN, FONT_EN)

    # ── 3c. Embed Images ──
    print('  - Embedding images...')
    _embed_images(doc, IMAGES_DIR)

    # ── 3d. Style Tables ──
    print('  - Styling tables...')
    _style_tables(doc, C_TABLE_HDR, BORDER_COLOR)

    # ── 3e. Set Document Metadata ──
    print('  - Setting metadata...')
    cp = doc.core_properties
    cp.title = '冀北风光储数智化生产支撑与数字孪生解决方案'
    cp.subject = '新能源场站数智化生产支撑系统技术方案书'
    cp.author = 'ZSSolutions'
    cp.language = 'zh-CN'
    cp.keywords = '冀北风光储;数智化;数字孪生;USDF;构网型;GOOSE;新能源场站'
    cp.category = '技术方案书'

    doc.save(str(OUTPUT_DOCX))
    print(f'  → Saved: {OUTPUT_DOCX}')
    print(f'  → Final size: {os.path.getsize(OUTPUT_DOCX):,} bytes')
    return True


def _add_cover_page(doc, c_primary, c_secondary, c_body, c_gray, font_cn, font_h):
    """Insert a professional cover page at the beginning using OOXML manipulation."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    from lxml import etree

    def _make_run_element(text, cn_font, en_font='Times New Roman', size=12,
                          bold=False, color=None, alignment=None):
        """Create a w:r element with font settings."""
        r = parse_xml(f'<w:r {nsdecls("w")}></w:r>')
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{cn_font}" w:ascii="{en_font}" w:hAnsi="{en_font}"/>')
        rPr.append(rFonts)
        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="{size * 2}"/>')
        rPr.append(sz)
        if bold:
            b = parse_xml(f'<w:b {nsdecls("w")}/>')
            rPr.append(b)
        if color:
            c = parse_xml(f'<w:color {nsdecls("w")} w:val="{color}"/>')
            rPr.append(c)
        r.insert(0, rPr)
        t = parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve">{text}</w:t>')
        r.append(t)
        return r

    def _make_para_element(alignment=None, space_before=0, space_after=0):
        """Create a w:p element with paragraph properties."""
        p = parse_xml(f'<w:p {nsdecls("w")}></w:p>')
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        if alignment is not None:
            jc = parse_xml(f'<w:jc {nsdecls("w")} w:val="{alignment}"/>')
            pPr.append(jc)
        if space_before:
            sb = parse_xml(f'<w:spacing {nsdecls("w")} w:before="{space_before}" w:after="{space_after}"/>')
            pPr.append(sb)
        p.insert(0, pPr)
        return p

    def _rgb_hex(color):
        """Convert RGBColor to hex string."""
        return str(color)  # RGBColor.__str__ returns hex like '1B3A5C'

    body = doc.element.body

    # Cover content: [spacers] + [项目全称] + [标题] + [客户] + [公司] + [日期] + [版本] + [page break]
    cover_items = [
        # (text, cn_font, size_pt, bold, hex_color, alignment, space_before_twips)
        ('冀北风光储数智化生产支撑与数字孪生系统', font_cn, 14, False, _rgb_hex(c_gray), 'center', 2400),
        ('技术方案书', font_h, 28, True, _rgb_hex(c_primary), 'center', 200),
        ('', font_cn, 10, False, _rgb_hex(c_body), 'center', 200),
        ('客户：冀北风光储基地', font_cn, 12, False, _rgb_hex(c_body), 'center', 1600),
        ('ZSSolutions', font_cn, 16, False, _rgb_hex(c_secondary), 'center', 600),
        ('2026年6月  |  版本 V2.0', font_cn, 10, False, _rgb_hex(c_gray), 'center', 1200),
    ]

    insert_pos = 0
    for text, cn_font, size, bold, hex_color, align, space_before in cover_items:
        p = _make_para_element(alignment=align, space_before=space_before, space_after=0)
        if text:
            r = _make_run_element(text, cn_font, size=size, bold=bold, color=hex_color)
            p.append(r)
        body.insert(insert_pos, p)
        insert_pos += 1

    # Add a page break after cover (as a separate empty paragraph with page break)
    pb_para = _make_para_element()
    pb_run = parse_xml(f'<w:r {nsdecls("w")}><w:br w:type="page"/></w:r>')
    pb_para.append(pb_run)
    body.insert(insert_pos, pb_para)

    # Now add a section break so cover section can have different header/footer
    # We add sectPr to the cover's last paragraph
    # This creates: Section 1 (cover, no header) | Section 2 (body, with header/footer)
    sect_pr = parse_xml(
        f'<w:sectPr {nsdecls("w")}>'
        f'  <w:type w:val="nextPage"/>'
        f'  <w:pgSz w:w="11906" w:h="16838"/>'
        f'</w:sectPr>'
    )
    body[insert_pos - 1].append(sect_pr)

    print('    Cover page inserted')


def _configure_headers_footers(doc, c_gray, font_cn, font_en):
    """Add headers and footers to body sections."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    def set_run_font(run, cn, en, size, color):
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = en
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), cn)

    for section in doc.sections:
        # ── Header ──
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            hp.paragraph_format.space_before = Pt(0)
            hp.paragraph_format.space_after = Pt(0)

            # Clean existing content and add header text
            for r in hp.runs:
                r.text = ''
            run = hp.add_run('冀北风光储数智化生产支撑系统  技术方案书')
            set_run_font(run, font_cn, font_en, 9, c_gray)

            # Add right-aligned company name via tab
            run2 = hp.add_run('\t\tZSSolutions')
            set_run_font(run2, font_cn, font_en, 9, c_gray)
        else:
            hp = header.add_paragraph()
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = hp.add_run('冀北风光储数智化生产支撑系统  技术方案书')
            set_run_font(run, font_cn, font_en, 9, c_gray)
            run2 = hp.add_run('\t\tZSSolutions')
            set_run_font(run2, font_cn, font_en, 9, c_gray)

        # ── Footer (page number) ──
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp.paragraph_format.space_before = Pt(0)
            fp.paragraph_format.space_after = Pt(0)
            # Clear existing
            for r in fp.runs:
                r.text = ''
            # Add page number field
            run = fp.add_run('- ')
            set_run_font(run, font_cn, font_en, 9, c_gray)
            # PAGE field
            r1 = fp.add_run()
            set_run_font(r1, font_cn, font_en, 9, c_gray)
            r1._element.append(parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
            r2 = fp.add_run()
            set_run_font(r2, font_cn, font_en, 9, c_gray)
            r2._element.append(parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
            r3 = fp.add_run()
            set_run_font(r3, font_cn, font_en, 9, c_gray)
            r3._element.append(parse_xml(
                f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
            run_end = fp.add_run(' -')
            set_run_font(run_end, font_cn, font_en, 9, c_gray)
        else:
            fp = footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run('- 第 PAGE 页 -')
            set_run_font(run, font_cn, font_en, 9, c_gray)

    print('    Headers/footers configured for all sections')


def _embed_images(doc, images_dir):
    """Replace figure caption paragraphs with actual embedded images + caption.

    Pandoc strips the image path when it can't find the file,
    leaving only the alt text as caption. We match by figure number.
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    # Map figure numbers to image files
    fig_to_image = {
        '图 1-1': 'fig1-1.png',
        '图 2-1': 'fig2-1.png',
        '图 4-0': 'fig4-0.png',
        '图 5-1': 'fig5-1.png',
        '图 6-1': 'fig6-1.png',
        '图 7-1': 'fig7-1.png',
        '图 8-1': 'fig8-1.png',
    }

    embedded = 0
    replaced_paras = set()

    for i, para in enumerate(doc.paragraphs):
        if i in replaced_paras:
            continue

        text = para.text.strip()
        for fig_id, img_filename in fig_to_image.items():
            if text.startswith(fig_id) and i not in replaced_paras:
                img_full_path = images_dir / img_filename
                if img_full_path.exists():
                    # Replace this paragraph's text with an image
                    # Clear existing runs
                    for run in para.runs:
                        run.text = ''

                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.space_before = Pt(12)
                    para.paragraph_format.space_after = Pt(6)

                    run = para.add_run()
                    try:
                        run.add_picture(str(img_full_path), width=Cm(14))
                        embedded += 1
                        replaced_paras.add(i)
                        # Also mark the next paragraph (blockquote caption) as handled
                        # Pandoc creates 3 paras: Body Text + 2x Block Text for fig caption
                        for offset in range(1, 4):
                            if i + offset < len(doc.paragraphs):
                                next_text = doc.paragraphs[i + offset].text.strip()
                                if next_text.startswith(fig_id):
                                    replaced_paras.add(i + offset)
                        print(f'    Embedded: {img_filename} at para {i}')
                    except Exception as e:
                        print(f'    WARNING: Failed to embed {img_filename}: {e}')
                else:
                    print(f'    WARNING: Image not found: {img_full_path}')
                break

    print(f'    Embedded {embedded} images')


def _style_tables(doc, header_bg_color, border_color):
    """Apply professional styling to all tables."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    C_BODY = RGBColor(0x33, 0x33, 0x33)
    C_ALT_ROW = RGBColor(0xF5, 0xF8, 0xFC)  # Very light blue for alternating rows

    styled = 0
    for table in doc.tables:
        if not table.rows:
            continue

        # Style header row
        header_row = table.rows[0]
        for cell in header_row.cells:
            # Set cell shading
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{header_bg_color}" w:val="clear"/>'
            )
            cell._element.get_or_add_tcPr().append(shading)

            # Set font for all paragraphs in header cell
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = C_BODY

        # Style data rows (alternating)
        for i, row in enumerate(table.rows[1:], 1):
            if i % 2 == 0:
                for cell in row.cells:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F5F8FC" w:val="clear"/>'
                    )
                    cell._element.get_or_add_tcPr().append(shading)

            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

        # Set table borders via XML
        tbl = table._element
        # Find or create tblPr
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        tblPr = tbl.find('.//w:tblPr', nsmap)
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
            tbl.insert(0, tblPr)

        # Set borders
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        styled += 1

    print(f'    Styled {styled} tables')


# ══════════════════════════════════════════════
#  Stage 4: Quality Verification
# ══════════════════════════════════════════════

def verify_docx():
    """Run quality checks on the final DOCX."""
    print('\n[4/4] Verifying DOCX quality...')

    try:
        from docx import Document
    except ImportError:
        print('  ERROR: python-docx not installed')
        return

    doc = Document(str(OUTPUT_DOCX))
    checks = []

    # Check 1: File size > 100KB
    size = os.path.getsize(OUTPUT_DOCX)
    checks.append(('文件大小 > 100KB', size > 100_000, f'{size/1024:.1f} KB'))

    # Check 2: Headers non-empty
    has_header = False
    for section in doc.sections:
        if section.header.paragraphs:
            for p in section.header.paragraphs:
                if p.text.strip():
                    has_header = True
                    break
        if has_header:
            break
    checks.append(('页眉非空', has_header, ''))

    # Check 3: Footer contains PAGE field
    has_page_field = False
    for section in doc.sections:
        footer = section.footer
        for p in footer.paragraphs:
            for run in p.runs:
                # Check for fldChar in run XML
                xml = run._element.xml
                if 'fldChar' in xml or 'PAGE' in run.text:
                    has_page_field = True
                    break
    checks.append(('页脚含页码', has_page_field, ''))

    # Check 4: §6.4 team capability section
    has_team_section = False
    for p in doc.paragraphs:
        if '团队能力背书' in p.text and p.style.name.startswith('Heading'):
            has_team_section = True
            break
    checks.append(('§6.4 团队能力背书', has_team_section, ''))

    # Check 5: Images embedded (check for drawings)
    drawing_count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            # Check for drawing elements
            for child in run._element:
                if child.tag.endswith('}drawing'):
                    drawing_count += 1
    checks.append(('图片已嵌入 (≥7)', drawing_count >= 7, f'{drawing_count} images'))

    # Check 6: Image files exist
    image_files = list(IMAGES_DIR.glob('fig*.png'))
    checks.append(('PNG图像文件存在', len(image_files) >= 7, f'{len(image_files)} files'))

    # Check 7: No residual ASCII box-drawing chars
    box_chars = set('┌┐└┘├┤┬┴┼│─')
    has_ascii = False
    for p in doc.paragraphs:
        if any(c in p.text for c in box_chars):
            has_ascii = True
            break
    checks.append(('无残留ASCII图', not has_ascii, ''))

    # Check 8: Tables have bold header
    tables_with_bold_header = 0
    for table in doc.tables:
        if table.rows:
            header_cell = table.rows[0].cells[0]
            for p in header_cell.paragraphs:
                for run in p.runs:
                    if run.bold:
                        tables_with_bold_header += 1
                        break
                else:
                    continue
                break
    checks.append(('表头粗体', tables_with_bold_header >= len(doc.tables) * 0.8,
                  f'{tables_with_bold_header}/{len(doc.tables)}'))

    # Check 9: TOC present
    has_toc = False
    for p in doc.paragraphs:
        if '目录' in p.text or 'Table of Contents' in p.text:
            has_toc = True
            break
    checks.append(('目录页存在', has_toc, ''))

    # Check 10: Section count >= 2 (cover + body, ideally)
    checks.append(('多节 (≥2)', len(doc.sections) >= 2, f'{len(doc.sections)} sections'))

    # Print results
    print('\n  ' + '=' * 60)
    print(f'  {"检查项":<30} {"结果":<10} {"详情"}')
    print('  ' + '-' * 60)
    passed = 0
    failed = 0
    for name, ok, detail in checks:
        status = '[PASS]' if ok else '[FAIL]'
        if ok:
            passed += 1
        else:
            failed += 1
        print(f'  {name:<30} {status:<10} {detail}')
    print('  ' + '-' * 60)
    print(f'  PASS: {passed}/{len(checks)}  FAIL: {failed}/{len(checks)}')
    print('  ' + '=' * 60)

    if failed > 0:
        print(f'\n  [WARN] {failed} checks failed. Review issues above.')
        if failed <= 2:
            print('  (Expected: residual ASCII in non-figure code blocks + section layout)')
    else:
        print(f'\n  [OK] All checks passed! DOCX ready for delivery.')

    return passed, failed


# ══════════════════════════════════════════════
#  Main
# ══════════════════════════════════════════════

def main():
    print('=' * 60)
    print('  JBFG Proposal DOCX Builder')
    print('=' * 60)

    check_only = '--check' in sys.argv

    if check_only:
        if not OUTPUT_DOCX.exists():
            print(f'ERROR: {OUTPUT_DOCX} not found')
            sys.exit(1)
        verify_docx()
        return

    # Stage 1: Preprocess MD
    preprocess_md()

    # Stage 2: Pandoc conversion
    if not run_pandoc():
        print('Build failed at Pandoc stage.')
        sys.exit(1)

    # Stage 3: Post-process DOCX
    if not postprocess_docx():
        print('Build failed at post-processing stage.')
        sys.exit(1)

    # Stage 4: Verify
    passed, failed = verify_docx()

    print(f'\nBuild complete: {OUTPUT_DOCX}')
    if failed > 0:
        sys.exit(1)


if __name__ == '__main__':
    main()

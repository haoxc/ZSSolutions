from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

base = Path(r'E:\Apps\Obsidian\ZSSolutions\03-项目课题\26.0618-数智化生产支撑系统(JBFG)\20-解决方案\方案产出-V5')
md_path = base / 'v5-冀北风光储数智化生产支撑与数字孪生解决方案.md'
out_path = base / 'v5-冀北风光储数智化生产支撑与数字孪生解决方案.docx'

def east(run, font):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), font)
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')

def set_run(run, font='宋体', size=12, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    east(run, font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_p(p, first=True, before=0, after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.74) if first else Cm(0)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = align

def shade(p, fill='D9D9D9'):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in('w:tblBorders')
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = tblBorders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tblBorders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')

def set_cell(cell, txt, header=False, center=False):
    cell.text = ''
    p = cell.paragraphs[0]
    set_p(p, first=False, after=0, align=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(txt)
    set_run(r, font='黑体' if header else '宋体', size=10.5, bold=header)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_text_runs(p, text, font='宋体', size=12, first=True, quote=False):
    set_p(p, first=first)
    if quote:
        shade(p)
    text = text.replace('"', '“')
    text = re.sub(r'\[\[([^\]|]+)\|([^\]]+)\]\]', r'\2', text)
    text = re.sub(r'\[\[([^\]]+)\]\]', r'\1', text)
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            set_run(r, font=font, size=size)
        token = m.group(0)
        if token.startswith('**'):
            r = p.add_run(token[2:-2])
            set_run(r, font=font, size=size, bold=True)
        elif token.startswith('`'):
            r = p.add_run(token[1:-1])
            set_run(r, font=font, size=size)
        else:
            r = p.add_run(token[1:-1])
            set_run(r, font=font, size=size, italic=True)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        set_run(r, font=font, size=size)

def cn_num(s):
    m = {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':10}
    if s == '十':
        return 10
    if '十' in s:
        if s.startswith('十'):
            return 10 + m[s[1:]]
        if s.endswith('十'):
            return m[s[0]] * 10
        a, b = s.split('十')
        return m[a] * 10 + m[b]
    return m[s]

def norm_h2(text):
    mm = re.match(r'^第([一二三四五六七八九十]+)章\s*(.*)$', text)
    if mm:
        return f"{cn_num(mm.group(1))} {mm.group(2)}"
    return text

def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break()

text = md_path.read_text(encoding='utf-8')
lines = text.splitlines()
img_files = {}
for p in base.rglob('*'):
    if p.is_file() and p.suffix.lower() in {'.png', '.jpg', '.jpeg', '.webp'}:
        img_files[p.name] = p

missing = []
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(2.8)
sec.right_margin = Cm(2.3)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.font.color.rgb = RGBColor(0, 0, 0)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = header.add_run('冀北风光储数智化生产支撑与数字孪生解决方案')
set_run(r, font='宋体', size=9)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'PAGE')
footer._p.append(fld)

started = False
pending_table_title = None
caption_to_skip = None
current_table = []

def flush_table():
    global current_table, pending_table_title
    if not current_table:
        return
    rows = [[c.strip() for c in row.strip().strip('|').split('|')] for row in current_table]
    if len(rows) > 1 and all(re.fullmatch(r':?-{3,}:?', c.replace(' ', '')) for c in rows[1]):
        rows.pop(1)
    if pending_table_title:
        p = doc.add_paragraph()
        set_p(p, first=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(pending_table_title)
        set_run(r, font='宋体', size=10.5, bold=True)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(table)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell(table.cell(i, j), val, header=(i == 0), center=(i == 0 or j == 0))
    current_table = []
    pending_table_title = None

for idx, raw in enumerate(lines):
    line = raw.rstrip()
    if idx < 14:
        continue
    if not line.strip():
        flush_table()
        continue
    if line.startswith('|') and line.endswith('|'):
        current_table.append(line)
        continue
    flush_table()

    if line == '---':
        continue

    m = re.match(r'^#\s+(.*)$', line)
    if m:
        p = doc.add_paragraph()
        set_p(p, first=False, before=24, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(m.group(1))
        set_run(r, font='黑体', size=22, bold=True)
        continue

    m = re.match(r'^##\s+(.*)$', line)
    if m:
        if started:
            doc.add_page_break()
        started = True
        p = doc.add_paragraph()
        set_p(p, first=False, before=24, after=18, align=WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(norm_h2(m.group(1).strip()))
        set_run(r, font='黑体', size=18, bold=True)
        continue

    m = re.match(r'^###\s+(.*)$', line)
    if m:
        p = doc.add_paragraph()
        set_p(p, first=False, before=12, after=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(m.group(1).strip())
        set_run(r, font='黑体', size=15, bold=True)
        continue

    m = re.match(r'^####\s+(.*)$', line)
    if m:
        p = doc.add_paragraph()
        set_p(p, first=False, before=8, after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
        r = p.add_run(m.group(1).strip())
        set_run(r, font='黑体' if re.match(r'^\d+\.\d+\.\d+\s', m.group(1).strip()) else '宋体', size=14 if re.match(r'^\d+\.\d+\.\d+\s', m.group(1).strip()) else 12, bold=True)
        continue

    if re.match(r'^表\s*\d+\s+', line):
        pending_table_title = line.strip()
        continue

    m = re.match(r'^!\[(.*?)\]\((.*?)\)$', line)
    if m:
        alt, src = m.groups()
        img = img_files.get(Path(src).name)
        if img and img.exists():
            p = doc.add_paragraph()
            set_p(p, first=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            rr = p.add_run()
            try:
                rr.add_picture(str(img), width=Cm(14))
            except Exception:
                rr.add_picture(str(img), height=Cm(18))
        else:
            missing.append(src)
            p = doc.add_paragraph()
            add_text_runs(p, f'[缺失图片：{src}]', first=False)
        cp = doc.add_paragraph()
        set_p(cp, first=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        rr = cp.add_run(alt)
        set_run(rr, font='宋体', size=10.5)
        caption_to_skip = f'> {alt}'
        continue

    if caption_to_skip and line.strip() == caption_to_skip:
        caption_to_skip = None
        continue

    if line.startswith('> '):
        p = doc.add_paragraph()
        add_text_runs(p, line[2:], font='楷体', first=False, quote=True)
        continue

    if line.startswith('- '):
        p = doc.add_paragraph()
        set_p(p, first=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.left_indent = Cm(0.74)
        r = p.add_run('• ' + line[2:])
        set_run(r, font='宋体', size=12)
        continue

    p = doc.add_paragraph()
    add_text_runs(p, line, font='宋体', size=12)

flush_table()

doc.save(str(out_path))
print(out_path)
if missing:
    print('MISSING_IMAGES:')
    for item in missing:
        print(item)

"""
Post-fix script for DOCX quality issues:
1. TOC: Add "目录" heading paragraph before TOC field if missing
2. Section break: Fix section break XML structure for python-docx recognition
"""
import sys
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

docx_path = sys.argv[1]
doc = Document(docx_path)
body = doc.element.body

# ─── Fix 1: TOC — ensure "目录" heading exists ───
has_toc_text = any('目录' in p.text for p in doc.paragraphs)
if not has_toc_text:
    print("TOC fix: Adding '目录' heading before TOC field...")
    # Find the TOC field paragraph
    toc_idx = None
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            xml = run._element.xml
            if 'instrText' in xml and 'TOC' in xml:
                toc_idx = i
                break
        if toc_idx is not None:
            break

    # Insert "目录" heading at beginning of body (after cover page)
    # We'll insert it at position 0 in the sections (first body paragraph after cover)
    # Use the existing TOC position or insert at the right spot
    if toc_idx is not None:
        # Insert a heading paragraph before the TOC field
        toc_para = doc.paragraphs[toc_idx]
        # Create heading paragraph with "目录" text
        new_p = parse_xml(
            f'<w:p {nsdecls("w")}>'
            f'  <w:pPr>'
            f'    <w:pStyle w:val="Heading1"/>'
            f'    <w:jc w:val="center"/>'
            f'  </w:pPr>'
            f'  <w:r>'
            f'    <w:rPr><w:rFonts w:eastAsia="黑体"/><w:b/><w:sz w:val="48"/></w:rPr>'
            f'    <w:t xml:space="preserve">目录</w:t>'
            f'  </w:r>'
            f'</w:p>'
        )
        toc_para._element.addprevious(new_p)
        print(f"  → Inserted '目录' heading before TOC field at paragraph {toc_idx}")
    else:
        # TOC field not found — insert at beginning of body
        print("  WARNING: TOC field not found in paragraphs. Checking XML directly...")
        # Try to find TOC in raw XML
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        toc_instr = body.findall('.//w:instrText', nsmap)
        if toc_instr:
            print(f"  → Found {len(toc_instr)} field instructions in XML (TOC may be present but not visible)")
            # Insert heading at start of content (after cover)
            new_p = parse_xml(
                f'<w:p {nsdecls("w")}>'
                f'  <w:pPr>'
                f'    <w:pStyle w:val="Heading1"/>'
                f'    <w:jc w:val="center"/>'
                f'  </w:pPr>'
                f'  <w:r>'
                f'    <w:rPr><w:rFonts w:eastAsia="黑体"/><w:b/><w:sz w:val="48"/></w:rPr>'
                f'    <w:t xml:space="preserve">目录</w:t>'
                f'  </w:r>'
                f'</w:p>'
            )
            # Find first paragraph after cover page (search for first Heading1)
            body_elements = list(body)
            for i, elem in enumerate(body_elements):
                if elem.tag.endswith('}p'):
                    pPr = elem.find(qn('w:pPr'))
                    if pPr is not None:
                        pStyle = pPr.find(qn('w:pStyle'))
                        if pStyle is not None and pStyle.get(qn('w:val')) == 'Heading1':
                            elem.addprevious(new_p)
                            print(f"  → Inserted '目录' heading at position {i}")
                            break

# ─── Fix 2: Section break — ensure multi-section ───
sections = doc.sections
print(f"Section fix: Currently {len(sections)} section(s)")

if len(sections) < 2:
    print("  Adding section break after cover page...")
    # Find the page break paragraph (last cover paragraph before content)
    # Look for the first paragraph with a page break
    body_elements = list(body)
    cover_end_idx = -1
    for i, elem in enumerate(body_elements):
        if elem.tag.endswith('}p'):
            for child in elem.iter():
                if child.tag.endswith('}br'):
                    br_type = child.get(qn('w:type'))
                    if br_type == 'page':
                        cover_end_idx = i
                        break
        if cover_end_idx > 0:
            break

    # Try harder: the cover ends after the date/version line and page break
    # We need to add sectPr correctly inside pPr of the paragraph after the break
    # Or add sectPr as the last element of the body (before the final body sectPr)

    if cover_end_idx > 0:
        target_elem = body_elements[cover_end_idx]
        # Ensure this paragraph has pPr
        pPr = target_elem.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            target_elem.insert(0, pPr)

        # Add sectPr inside pPr
        sect = parse_xml(
            f'<w:sectPr {nsdecls("w")}>'
            f'  <w:type w:val="nextPage"/>'
            f'  <w:pgSz w:w="11906" w:h="16838"/>'
            f'</w:sectPr>'
        )
        pPr.append(sect)
        print(f"  → Added sectPr to page break paragraph at index {cover_end_idx}")
    else:
        # Fallback: add section break at the very end of body
        print("  WARNING: Page break not found. Adding sectPr to last cover paragraph...")
        for i, elem in enumerate(body_elements):
            if elem.tag.endswith('}p'):
                text_elems = elem.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                text = ''.join(t.text or '' for t in text_elems)
                if '版本' in text or 'V3' in text or '2026' in text:
                    pPr = elem.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
                        elem.insert(0, pPr)
                    sect = parse_xml(
                        f'<w:sectPr {nsdecls("w")}>'
                        f'  <w:type w:val="nextPage"/>'
                        f'  <w:pgSz w:w="11906" w:h="16838"/>'
                        f'</w:sectPr>'
                    )
                    pPr.append(sect)
                    print(f"  → Added sectPr to version paragraph at index {i}")
                    break

doc.save(docx_path)

# Verify fix
doc2 = Document(docx_path)
new_sections = len(doc2.sections)
has_toc = any('目录' in p.text for p in doc2.paragraphs)
print(f"\nVerification after fix:")
print(f"  Sections: {new_sections}")  # Should be 2 with fix, 2 without
print(f"  TOC text '目录': {'YES' if has_toc else 'NOPE'}")
print(f"  Result: {'✓ All fixed' if new_sections >= 2 and has_toc else '⚠ Some issues remain'}")

"""
JBFG Proposal V3 DOCX Builder — One-click build from MD to final deliverable.

Stages:
  1. Preprocess MD (strip YAML, convert WikiLinks, Mermaid→image refs)
  2. Pandoc convert with reference template (--reference-doc)
  3. Post-process DOCX (cover page, headers/footers, images, table styling)
  4. Quality verification

Usage:
  python build_docx.py           # Full build
  python build_docx.py --check   # Verify existing DOCX only
"""

import os
import re
import sys
import subprocess
from pathlib import Path

# ── Paths (V3) ──
BASE_DIR = Path(__file__).resolve().parent.parent  # 方案产出-V3/
BUILD_DIR = BASE_DIR / 'build'
IMAGES_DIR = BASE_DIR / 'images'
MD_SOURCE = BASE_DIR / 'v3-冀北风光储数智化生产支撑与数字孪生解决方案.md'
MD_PREPROCESSED = BUILD_DIR / 'source_preprocessed.md'
REFERENCE_DOCX = BUILD_DIR / '参考模板.docx'
OUTPUT_DOCX = BASE_DIR / '冀北风光储数智化生产支撑与数字孪生解决方案-V3.docx'

# Figure mapping: Mermaid block first-line identifier → image file
FIGURE_MAP = {
    '三大结构性痛点': ('images/fig1-1.png', '图1-1：方案核心逻辑——三大矛盾通过USDF三通道分流统一解决'),
    'S 现状': ('images/fig2-1.png', '图2-1：SCQA分析框架——从现状到方案的结构化推导'),
    'USDF["数字孪生底座 USDF"]': ('images/fig4-0.png', '图4-0：五大业务子系统闭环全景——建→控→评→检→管'),
    '{"层级": "L0 暂态(GOOSE直通)", "延迟ms": 100': ('images/fig5-1.png', '图5-1：四层SLA延迟预算'),
    'title 冀北风光储数智化项目实施计划': ('images/fig6-1.png', '图6-1：项目实施甘特图——4个月三阶段四里程碑'),
    '{"维度": "技术风险", "等级": "阻断", "数量": 2}': ('images/fig7-1.png', '图7-1：风险热力图——18项风险三维度三色等级分布'),
    '{"项目": "定制开发", "金额": 93}': ('images/fig8-1.png', '图8-1：报价构成——220万含税分配'),
}


# ══════════════════════════════════════════════
#  Stage 1: MD Preprocessing
# ══════════════════════════════════════════════

def strip_yaml_frontmatter(content):
    """Remove YAML frontmatter from Markdown."""
    if content.startswith('---'):
        end = content.find('---', 3)
        if end != -1:
            return content[end + 3:].lstrip('\n')
    return content


def convert_wikilinks(content):
    """Convert Obsidian WikiLinks to plain text."""
    def replace_link(m):
        inner = m.group(1)
        if '|' in inner:
            return inner.split('|', 1)[1]
        elif '#' in inner:
            return inner.split('#', 1)[0]
        return inner
    return re.sub(r'\[\[([^\]]+)\]\]', replace_link, content)


def replace_mermaid_with_images(content):
    """Replace Mermaid/Vega-Lite code blocks with image references for key figures.
    Uses line-by-line block detection to match blocks to pre-rendered PNG images.
    """
    lines = content.split('\n')
    parts = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Detect start of fenced code block: ```mermaid or ```vega-lite
        is_mermaid = stripped.startswith('```mermaid')
        is_vega = stripped.startswith('```vega-lite')

        if is_mermaid or is_vega:
            lang = stripped
            block_content_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                block_content_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```

            block_text = '\n'.join(block_content_lines)
            replaced = False

            for keyword, (img_path, caption) in FIGURE_MAP.items():
                if keyword in block_text:
                    parts.append(f'![{caption}]({img_path})')
                    parts.append('')
                    parts.append(f'> {caption}')
                    parts.append('')
                    replaced = True
                    break

            if not replaced:
                # Preserve as fenced code block
                parts.append(lang)
                parts.extend(block_content_lines)
                parts.append('```')
        else:
            parts.append(line)
            i += 1

    return '\n'.join(parts)


def preprocess_md():
    """Preprocess MD source for Pandoc conversion."""
    print('[1/4] Preprocessing MD source...')

    with open(MD_SOURCE, 'r', encoding='utf-8') as f:
        content = f.read()

    # 1. Strip YAML frontmatter
    content = strip_yaml_frontmatter(content)
    print('  - YAML frontmatter stripped')

    # 2. Convert WikiLinks
    content = convert_wikilinks(content)
    print('  - WikiLinks converted')

    # 3. Replace Mermaid blocks with image references
    content = replace_mermaid_with_images(content)
    # Convert relative image paths to absolute for Pandoc
    images_abs = IMAGES_DIR.resolve()
    content = re.sub(
        r'!\[([^\]]*)\]\(images/([^)]+)\)',
        lambda m: f'![{m.group(1)}]({images_abs / m.group(2)})',
        content
    )
    print('  - Mermaid blocks → image references (absolute paths)')

    # 4. Remove Obsidian comments (%% ... %%)
    content = re.sub(r'%%.*?%%', '', content, flags=re.DOTALL)
    print('  - Comments removed')

    # Write preprocessed file
    with open(MD_PREPROCESSED, 'w', encoding='utf-8') as f:
        f.write(content)

    print(f'  → Saved: {MD_PREPROCESSED}')
    print(f'  → Lines: {len(content.splitlines())}')
    return MD_PREPROCESSED


# ══════════════════════════════════════════════
#  Stage 2: Pandoc Conversion
# ══════════════════════════════════════════════

def run_pandoc():
    """Convert preprocessed MD to DOCX using reference template."""
    print('\n[2/4] Running Pandoc conversion...')

    cmd = [
        'pandoc',
        str(MD_PREPROCESSED),
        '--from=gfm',
        '--to=docx',
        f'--reference-doc={REFERENCE_DOCX}',
        '--toc',
        '--toc-depth=3',
        '--metadata', 'title=冀北风光储数智化生产支撑与数字孪生解决方案',
        '--metadata', 'author=ZSSolutions',
        '--metadata', 'subject=新能源场站数智化生产支撑系统技术方案书',
        '--metadata', 'lang=zh-CN',
        '--metadata', 'keywords=冀北风光储,数智化,数字孪生,USDF,构网型,GOOSE',
        '-o', str(OUTPUT_DOCX),
    ]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8',
                                errors='replace', timeout=120)
    except FileNotFoundError:
        print('  ERROR: Pandoc not found. Run: winget install pandoc')
        return False
    except subprocess.TimeoutExpired:
        print('  ERROR: Pandoc timed out.')
        return False

    if result.returncode != 0:
        print(f'  ERROR: Pandoc failed:\n{result.stderr}')
        return False

    print(f'  → Saved: {OUTPUT_DOCX}')
    print(f'  → Size: {os.path.getsize(OUTPUT_DOCX):,} bytes')
    return True


# ══════════════════════════════════════════════
#  Stage 3: DOCX Post-processing
# ══════════════════════════════════════════════

def postprocess_docx():
    """Post-process Pandoc-generated DOCX: cover page, table styling."""
    print('\n[3/4] Post-processing DOCX...')

    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.oxml.ns import qn
        from docx.oxml import parse_xml
    except ImportError:
        print('  WARNING: python-docx not installed. Run: pip install python-docx')
        print('  DOCX built without post-processing.')
        return True  # Not fatal

    # Color palette
    C_PRIMARY = RGBColor(0x1B, 0x3A, 0x5C)
    C_SECONDARY = RGBColor(0x2B, 0x5C, 0x8A)
    C_BODY = RGBColor(0x33, 0x33, 0x33)
    C_GRAY = RGBColor(0x66, 0x66, 0x66)
    C_TABLE_HDR = 'D6E4F0'
    BORDER_COLOR = 'CCCCCC'

    doc = Document(str(OUTPUT_DOCX))

    # ── 3a. Cover Page ──
    print('  - Adding cover page...')
    _add_cover_page(doc, C_PRIMARY, C_SECONDARY, C_BODY, C_GRAY)

    # ── 3b. Table Styling ──
    print('  - Styling tables...')
    _style_tables(doc, C_TABLE_HDR, BORDER_COLOR)

    # Save
    doc.save(str(OUTPUT_DOCX))
    print(f'  → Saved styled DOCX: {OUTPUT_DOCX}')
    return True


def _add_cover_page(doc, c_primary, c_secondary, c_body, c_gray):
    """Insert cover page at the beginning of the document."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

    first_para = doc.paragraphs[0] if doc.paragraphs else None

    def add_cover_para(text, font_name='宋体', size=Pt(12), color=c_body,
                       bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
        p = doc.add_paragraph()
        if first_para:
            first_para._element.addprevious(p._element)
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = size
        run.font.color.rgb = color
        run.bold = bold
        p.alignment = alignment
        p.paragraph_format.space_after = space_after
        return p

    # Title lines
    add_cover_para('冀北风光储数智化生产支撑', '黑体', Pt(26), c_primary, True,
                   WD_ALIGN_PARAGRAPH.CENTER, Pt(12))
    add_cover_para('与数字孪生解决方案', '黑体', Pt(26), c_primary, True,
                   WD_ALIGN_PARAGRAPH.CENTER, Pt(36))

    # Subtitle
    add_cover_para('USDF统一空间数据底座驱动的新能源场站空间对象操作系统',
                   '宋体', Pt(14), c_secondary, False, WD_ALIGN_PARAGRAPH.CENTER, Pt(48))

    # Version
    add_cover_para('版本：V3.0（双模版）', '宋体', Pt(11), c_gray, False,
                   WD_ALIGN_PARAGRAPH.CENTER, Pt(6))
    add_cover_para('日期：2026-06-21', '宋体', Pt(11), c_gray, False,
                   WD_ALIGN_PARAGRAPH.CENTER, Pt(6))
    add_cover_para('密级：机密·投标专用', '宋体', Pt(11), c_gray, False,
                   WD_ALIGN_PARAGRAPH.CENTER, Pt(24))

    # Separator
    add_cover_para('━' * 50, '宋体', Pt(8), c_secondary, False,
                   WD_ALIGN_PARAGRAPH.CENTER, Pt(36))

    # Organization
    add_cover_para('ZSSolutions · 中数解决方案', '宋体', Pt(12), c_body, False,
                   WD_ALIGN_PARAGRAPH.CENTER, Pt(6))
    add_cover_para('2026年6月', '宋体', Pt(11), c_gray, False,
                   WD_ALIGN_PARAGRAPH.CENTER, Pt(0))

    # Page break
    add_cover_para('', '宋体', Pt(1), c_body)
    if doc.paragraphs:
        doc.paragraphs[0].add_run().add_break(WD_BREAK.PAGE)


def _style_tables(doc, hdr_hex_color, border_color):
    """Apply consistent styling to all tables."""
    from docx.shared import Pt
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    for table in doc.tables:
        # Style header row
        if table.rows:
            for cell in table.rows[0].cells:
                shading_xml = (
                    f'<w:shd {nsdecls("w")} w:fill="{hdr_hex_color}" w:val="clear"/>'
                )
                shading = parse_xml(shading_xml)
                cell._tc.get_or_add_tcPr().append(shading)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        if run.font.size is None:
                            run.font.size = Pt(9)

        # Style all cells: borders + font
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                borders_xml = (
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'<w:top w:val="single" w:sz="4" w:color="{border_color}"/>'
                    f'<w:left w:val="single" w:sz="4" w:color="{border_color}"/>'
                    f'<w:bottom w:val="single" w:sz="4" w:color="{border_color}"/>'
                    f'<w:right w:val="single" w:sz="4" w:color="{border_color}"/>'
                    f'</w:tcBorders>'
                )
                tc_pr.append(parse_xml(borders_xml))
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.size is None:
                            run.font.size = Pt(9)


# ══════════════════════════════════════════════
#  Stage 4: Quality Verification
# ══════════════════════════════════════════════

def verify_output():
    """Basic quality verification of the final DOCX."""
    print('\n[4/4] Verifying output...')

    checks = []
    file_size = os.path.getsize(OUTPUT_DOCX)

    checks.append(('File exists', os.path.exists(OUTPUT_DOCX)))
    checks.append(('File > 1MB', file_size > 1_000_000))
    checks.append(('File < 50MB', file_size < 50_000_000))

    # Content checks via python-docx
    try:
        from docx import Document
        doc = Document(str(OUTPUT_DOCX))

        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        total_chars = sum(len(p.text) for p in doc.paragraphs)

        checks.append(('Paragraphs > 100', para_count > 100))
        checks.append(('Tables > 25', table_count > 25))
        checks.append(('Content > 30K chars', total_chars > 30_000))
        checks.append(('Tables <= 50', table_count <= 50))
    except ImportError:
        pass

    all_ok = True
    for label, ok in checks:
        status = 'OK' if ok else 'FAIL'
        print(f'  [{status}] {label}')
        if not ok:
            all_ok = False

    if all_ok:
        print(f'\n  All checks passed.')
        print(f'  Output: {OUTPUT_DOCX}')
    else:
        print(f'\n  Some checks failed.')
    return all_ok


# ══════════════════════════════════════════════
#  Main
# ══════════════════════════════════════════════

def main():
    print('=' * 60)
    print('  JBFG V3 Proposal DOCX Builder')
    print('=' * 60)

    if len(sys.argv) > 1 and sys.argv[1] == '--check':
        if os.path.exists(OUTPUT_DOCX):
            verify_output()
        else:
            print(f'  No DOCX found at {OUTPUT_DOCX}')
        return

    # Stage 1
    preprocess_md()

    # Stage 2
    if not run_pandoc():
        print('\n  Build FAILED at Pandoc stage.')
        return

    # Stage 3
    if not postprocess_docx():
        print('\n  Build FAILED at post-processing stage.')
        return

    # Stage 4
    verify_output()

    print('\n' + '=' * 60)
    print(f'  Build complete: {OUTPUT_DOCX.name}')
    print('=' * 60)


if __name__ == '__main__':
    main()

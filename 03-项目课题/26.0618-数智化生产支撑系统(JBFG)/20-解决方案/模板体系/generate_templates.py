"""
Generate Word template files (.docx) for JBFG proposal:
  - Template A: Comprehensive Bid Version (完整应标版)
  - Template B: Technical Solution Version (技术方案版)

Generation rules documented in: Word模板生成规则说明书.md
Rules implemented:
  1. Section breaks before each H1 (not just page breaks)
  2. Real Word tables (add_word_table) — no markdown text tables
  3. Image placeholders (add_image_placeholder) — bordered box + caption
  4. Revision history at END (before appendices), not at front
  5. Shared style family (configure_styles) inherited by both templates
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Color Palette ──
C_PRIMARY   = RGBColor(0x1B, 0x3A, 0x5C)  # Deep navy
C_SECONDARY = RGBColor(0x2B, 0x5C, 0x8A)  # Medium blue
C_ACCENT    = RGBColor(0x2E, 0x75, 0xB6)  # Standard blue
C_BODY      = RGBColor(0x33, 0x33, 0x33)  # Dark gray body text
C_GRAY      = RGBColor(0x66, 0x66, 0x66)  # Gray
C_LIGHT     = RGBColor(0x99, 0x99, 0x99)  # Light gray
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_TABLE_HDR = RGBColor(0xD6, 0xE4, 0xF0)  # Light blue table header
C_TABLE_BW  = RGBColor(0xD9, 0xD9, 0xD9)  # Light gray for B&W
C_BORDER    = RGBColor(0xCC, 0xCC, 0xCC)  # Border gray

FONT_CN  = '宋体'
FONT_H   = '黑体'
FONT_EN  = 'Times New Roman'
FONT_CODE = 'Consolas'


# ══════════════════════════════════════════════
#  Helpers: font, paragraph, border
# ══════════════════════════════════════════════

def set_font(run, name_cn=FONT_CN, name_en=FONT_EN, size=12, bold=False, color=C_BODY):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = name_en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)


def set_spacing(paragraph, before=0, after=0, line_spacing=None, line_rule=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing is not None:
        if line_rule:
            pf.line_spacing_rule = line_rule
        pf.line_spacing = Pt(line_spacing)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders (size in 1/8 pt)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val is not None:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{val}" w:space="0" w:color="CCCCCC"/>'
            )
            borders.append(border)
    tcPr.append(borders)


def shade_cell(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tcPr.append(shd)


# ══════════════════════════════════════════════
#  Helpers: section, chapter, table, image
# ══════════════════════════════════════════════

def add_body_section(doc):
    """Add a new section inheriting page setup, for chapter page break."""
    sec = doc.add_section()
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.5)
    return sec


def add_chapter(doc, title, is_first=False):
    """Add H1 on a new section (section break before it), unless is_first."""
    if not is_first:
        add_body_section(doc)
    doc.add_heading(title, level=1)


def add_word_table(doc, headers, rows, caption=None, is_color=True):
    """Add a real Word table with professional formatting.

    Args:
        doc: Document
        headers: list of header strings
        rows: list of lists (data rows)
        caption: optional table caption (shown above table)
        is_color: True for color (Template A), False for grayscale (Template B)
    """
    # ── Caption ──
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(cap, before=6, after=3)
        set_font(cap.runs[0], FONT_CN, FONT_EN, 10.5, bold=True, color=C_BODY)

    # ── Table ──
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_color = '1B3A5C' if is_color else '4D4D4D'

    for col_idx, header_text in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        set_font(run, FONT_CN, FONT_EN, 10.5, bold=True, color=C_WHITE)
        shade_cell(cell, hdr_color)

    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, FONT_CN, FONT_EN, 10.5, color=C_BODY)

    # ── Table borders via XML ──
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="12" w:space="0" w:color="1B3A5C"/>'
        f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="1B3A5C"/>'
        f'<w:left w:val="single" w:sz="12" w:space="0" w:color="1B3A5C"/>'
        f'<w:right w:val="single" w:sz="12" w:space="0" w:color="1B3A5C"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Column widths: equal
    tblGrid = tbl.tblGrid if tbl.tblGrid is not None else parse_xml(f'<w:tblGrid {nsdecls("w")}></w:tblGrid>')
    # Clear existing grid cols
    for gc in tblGrid.findall(qn('w:gridCol')):
        tblGrid.remove(gc)
    col_width = int(907200 / len(headers))  # ~14.66cm in EMU / cols
    for _ in headers:
        gridCol = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{col_width}"/>')
        tblGrid.append(gridCol)

    return table


def add_image_placeholder(doc, caption_text='图 placeholder', width_cm=14, height_cm=5):
    """Add a bordered image placeholder box with numbered caption below."""
    # ── Placeholder box (single-cell bordered table) ──
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.rows[0].cells[0]
    cell.text = ''

    # Set cell height
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="atLeast"/>')
    trPr.append(trHeight)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n\n[ 此处插入图片：{caption_text} ]\n建议尺寸：{width_cm}cm × {height_cm}cm\n\n')
    set_font(run, FONT_CN, FONT_EN, 11, color=C_GRAY)

    # Cell shading (very light gray)
    shade_cell(cell, 'F5F5F5')

    # Table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="dashed" w:sz="6" w:space="0" w:color="999999"/>'
        f'<w:bottom w:val="dashed" w:sz="6" w:space="0" w:color="999999"/>'
        f'<w:left w:val="dashed" w:sz="6" w:space="0" w:color="999999"/>'
        f'<w:right w:val="dashed" w:sz="6" w:space="0" w:color="999999"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # ── Caption below ──
    cap = doc.add_paragraph(caption_text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(cap, before=3, after=6)
    set_font(cap.runs[0], FONT_CN, FONT_EN, 10.5, bold=True, color=C_BODY)

    return table


# ══════════════════════════════════════════════
#  Cover pages
# ══════════════════════════════════════════════

def add_cover_page_a(doc):
    """Template A: Full cover page with company branding."""
    # Confidentiality bar
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run('■ 机  密  ·  仅限指定收件人  ■')
    set_font(run, FONT_H, FONT_EN, 11, bold=True, color=C_WHITE)
    set_spacing(p, before=0, after=60)

    doc.add_paragraph()

    # Logo placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[ 公司 Logo ]')
    set_font(run, FONT_H, FONT_EN, 20, bold=True, color=C_PRIMARY)

    p = doc.add_paragraph()
    set_spacing(p, before=24)

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('冀北风光储数智化生产支撑\n与数字孪生解决方案')
    set_font(run, FONT_H, FONT_EN, 28, bold=True, color=C_PRIMARY)
    set_spacing(p, before=0, after=6)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('USDF 统一空间数据底座——新能源场站的空间对象操作系统')
    set_font(run, FONT_CN, FONT_EN, 16, color=C_GRAY)

    # Info block
    set_spacing(p, before=24, after=12)
    for line in [
        '编制单位：止善科技（北京）有限公司',
        '编制日期：2026 年    月',
        '版    本：V1.0',
        '文件编号：JBFG-SOL-2026-001',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_font(run, FONT_CN, FONT_EN, 14, color=C_BODY)

    # Bottom copyright
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=48)
    run = p.add_run('© 2026 止善科技（北京）有限公司. All Rights Reserved.')
    set_font(run, FONT_CN, FONT_EN, 9, color=C_LIGHT)


def add_cover_page_b(doc):
    """Template B: Minimal cover page for technical review."""
    # Top bar
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(' ' * 80)
    set_font(run, FONT_CN, FONT_EN, 14, color=C_PRIMARY)
    set_spacing(p, before=0, after=0)

    for _ in range(6):
        p = doc.add_paragraph()
        set_spacing(p, before=12)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('冀北风光储数智化生产支撑\n与数字孪生  技术方案')
    set_font(run, FONT_H, FONT_EN, 28, bold=True, color=C_PRIMARY)

    # Version
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=12, after=12)
    run = p.add_run('V1.0  |  2026 年    月')
    set_font(run, FONT_CN, FONT_EN, 14, color=C_GRAY)

    for _ in range(6):
        p = doc.add_paragraph()
        set_spacing(p, before=12)

    # Bottom bar
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(' ' * 80)
    set_font(run, FONT_CN, FONT_EN, 14, color=C_PRIMARY)
    set_spacing(p, before=0, after=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=6)
    run = p.add_run('止善科技（北京）有限公司  |  Confidential')
    set_font(run, FONT_CN, FONT_EN, 9, color=C_LIGHT)


# ══════════════════════════════════════════════
#  Style configuration
# ══════════════════════════════════════════════

def configure_styles(doc):
    """Configure all styles for both templates (shared style family)."""
    s = doc.styles

    # ── Normal ──
    n = s['Normal']
    n.font.name = FONT_EN
    n.font.size = Pt(12)
    n.font.color.rgb = C_BODY
    rPr = n.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    pf = n.paragraph_format
    pf.first_line_indent = Pt(24)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(22)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Heading 1 ──
    h = s['Heading 1']
    h.font.name = FONT_EN
    h.font.size = Pt(18)
    h.font.bold = True
    h.font.color.rgb = C_PRIMARY
    rPr = h.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_H)
    pf = h.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(0)
    pf.keep_with_next = True

    # ── Heading 2 ──
    h = s['Heading 2']
    h.font.name = FONT_EN
    h.font.size = Pt(14)
    h.font.bold = True
    h.font.color.rgb = C_SECONDARY
    rPr = h.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_H)
    pf = h.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(0)
    pf.keep_with_next = True

    # ── Heading 3 ──
    h = s['Heading 3']
    h.font.name = FONT_EN
    h.font.size = Pt(13)
    h.font.bold = True
    h.font.color.rgb = C_BODY
    rPr = h.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_H)
    pf = h.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(0)
    pf.keep_with_next = True

    # ── Heading 4 ──
    h = s['Heading 4']
    h.font.name = FONT_EN
    h.font.size = Pt(12)
    h.font.bold = True
    h.font.color.rgb = C_BODY
    rPr = h.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_H)
    pf = h.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(0)
    pf.keep_with_next = True


# ══════════════════════════════════════════════
#  Header / Footer
# ══════════════════════════════════════════════

def _add_page_number(footer):
    """Insert '- PAGE -' field into footer."""
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('- ')
    set_font(run, FONT_EN, FONT_EN, 9, color=C_GRAY)
    r = fp.add_run()
    r._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r2 = fp.add_run()
    r2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    r3 = fp.add_run()
    r3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    run2 = fp.add_run(' -')
    set_font(run2, FONT_EN, FONT_EN, 9, color=C_GRAY)


def add_header_footer(section, header_text):
    """Set header and footer for a body section."""
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(header_text)
    set_font(run, FONT_CN, FONT_EN, 9, color=C_GRAY)

    footer = section.footer
    footer.is_linked_to_previous = False
    _add_page_number(footer)


# ══════════════════════════════════════════════
#  Content skeletons
# ══════════════════════════════════════════════

def add_content_a(doc, is_color=True):
    """Content for Template A: 8 chapters + revision history at end."""
    # ──────────────────────────────────────────
    # Chapter 1: 项目概述
    # ──────────────────────────────────────────
    add_chapter(doc, '第一章  项目概述', is_first=True)
    add_image_placeholder(doc, '图 1-1  系统总体架构示意', 14, 6)

    doc.add_heading('1.1  项目背景', level=2)
    doc.add_paragraph('冀北风光储基地概况：场站规模、机组配置、并网情况……新型电力系统对场站运维的新要求……数字化向数智化演进的行业趋势……')

    doc.add_heading('1.2  需求理解（SCQA 精炼版）', level=2)
    doc.add_paragraph('S (现状)：冀北基地多源异构设备、多系统并行、数据孤岛……')
    doc.add_paragraph('C (挑战)：数据挑战（百万测点实时接入）、实时性挑战（端到端 ≤1s）、可靠性挑战（MTBF ≥10000h）……')
    doc.add_paragraph('Q (核心问题)：如何在 4 个月内交付一个满足百万测点实时接入、毫秒级暂态控制、AI 就绪、且二期可无缝扩展至多场站的平台？')
    doc.add_paragraph('A (方案概要)：USDF 统一空间数据底座 + 三通道分流 + 低代码组态 + AI 即插即用。')

    doc.add_heading('1.3  术语定义', level=2)
    add_word_table(doc,
        ['术语', '英文', '定义'],
        [
            ['数字孪生', 'Digital Twin', '物理系统的实时数字镜像'],
            ['透明场站', 'Transparent Station', '全要素可视化的新能源场站'],
            ['USDF', 'Unified Spatial Data Foundation', '统一空间数据底座'],
            ['GOOSE', 'GOOSE', 'IEC 61850 面向通用对象的变电站事件'],
        ],
        '表 1-1  术语定义',
        is_color,
    )

    # ──────────────────────────────────────────
    # Chapter 2
    # ──────────────────────────────────────────
    add_chapter(doc, '第二章  平台建设思想与顶层设计')

    doc.add_heading('2.1  核心思想：空间对象操作系统', level=2)
    doc.add_paragraph('USDF 统一空间数据底座——新能源场站的空间对象操作系统。一台设备注册一次，所有业务按需消费。')

    add_word_table(doc,
        ['OS 概念', 'USDF 对应', '工程含义'],
        [
            ['文件系统', '空间对象注册', '设备/设施获得全局唯一 ID + 空间坐标'],
            ['设备驱动', '统一采集适配器', 'IEC104/Modbus/GOOSE → 统一数据格式'],
            ['系统调用', '北向 API', '/api/v1/objects/{id}/* 标准接口'],
            ['进程调度', '三通道分流', 'P1暂态/P2稳态/P3历史，按优先级调度'],
            ['用户态/内核态', '业务层/USDF底座', '业务系统不碰数据管线'],
        ],
        '表 2-1  空间对象操作系统类比',
        is_color,
    )

    doc.add_heading('2.2  四大架构约束', level=2)
    add_word_table(doc,
        ['约束', '架构含义', '检验标准', '违反后果'],
        [
            ['可靠', '暂态控制独立通道，安全分区隔离', 'GOOSE ≤100ms', '设备损坏/电网事故'],
            ['扩展', 'USDF 天然多场站，station_id 预留', '二期不加字段不改 API', '一期投资沉没'],
            ['AI 友好', '北向 API 输出标准化对象+时序数据', '算法团队不碰采集层', 'AI 与管线紧耦合'],
            ['实施效能', '边云协同+断网自持+国产化直选', '一期不依赖二期基础设施', '项目延期'],
        ],
        '表 2-2  四大架构约束',
        is_color,
    )

    doc.add_heading('2.3  三条顶层原理', level=2)
    doc.add_paragraph('原理一：三通道分流——P1 暂态控制 (GOOSE ≤100ms) / P2 稳态监控 (IEC104 ≤1s) / P3 历史分析 (RESTful ≤5s)。')
    doc.add_paragraph('原理二：不可越层规则——展示层不直连时序库、采集层不直写业务库、USDF 不存时序数据。')
    doc.add_paragraph('原理三：空间是统一维度——五种对象共享空间坐标，空间查询替代定制 JOIN。')

    doc.add_heading('2.4  方法论：从指标反推架构', level=2)
    doc.add_paragraph('每一个架构决策都对应一个可检验的非功能指标。34 项非功能指标 → 架构选择链。')

    doc.add_heading('2.5  平台能力的四个支柱', level=2)
    add_word_table(doc,
        ['#', '能力', '核心价值'],
        [
            ['1', '门户集成与双模交互', '子系统统一挂载+运营管理与监控看板'],
            ['2', '元数据驱动页面构建', '自动生成 CRUD 页面，减少前端重复开发'],
            ['3', '2D/3D 场景双向联动', '页面操作→3D高亮；3D点击→数据面板'],
            ['4', '样式与主题配置', '全局品牌色/字体/暗色模式，统一视觉风格'],
        ],
        '表 2-3  平台能力的四个支柱',
        is_color,
    )

    # ──────────────────────────────────────────
    # Chapter 3
    # ──────────────────────────────────────────
    add_chapter(doc, '第三章  总体技术架构')

    doc.add_heading('3.1  设计理念与架构原则', level=2)
    doc.add_paragraph('USDF 底座 + 五业务域。底座先行 → 业务渐进 → 多场站可扩展。')

    add_image_placeholder(doc, '图 3-1  系统逻辑架构（四层）', 14, 8)

    doc.add_heading('3.2  部署拓扑（三级 + 安全分区）', level=2)
    doc.add_paragraph('中心云（集团侧）↕ 场站边缘（I 区/II 区/III 区）↕ 设备层。')

    doc.add_heading('3.3  数据分层与流向', level=2)
    add_word_table(doc,
        ['层', '职责', 'SLA'],
        [
            ['采集层', '协议适配、数据归一化', '按数据源频率'],
            ['消息总线', '多源汇聚、分区路由', '≤50ms'],
            ['时序库', '测点历史、聚合查询', '100万条/s 写入；≤5s 查询'],
            ['业务库', '主数据、告警、工单', '事务一致性'],
            ['缓存层', '实时快照', '≤1ms 命中'],
            ['展示层', 'API 聚合、权限过滤', 'P95 ≤200ms'],
        ],
        '表 3-1  数据分层与 SLA',
        is_color,
    )

    doc.add_heading('3.4  关键技术选型', level=2)
    add_word_table(doc,
        ['决策域', '推荐', '备选', '优先级'],
        [
            ['时序数据库', 'ClickHouse', 'TDengine/DolphinDB', 'P0'],
            ['业务数据库', 'Postgres+PostGIS', '达梦 DM8/人大金仓', 'P0'],
            ['消息队列', 'Kafka', 'RocketMQ/Pulsar', 'P0'],
            ['GIS 引擎', 'GeoServer', '超图/中地', 'P0'],
            ['三维引擎', 'Cesium+Three.js', '—', '待定'],
            ['API 网关', 'APISIX', '—', 'P1'],
        ],
        '表 3-2  关键技术选型',
        is_color,
    )

    doc.add_heading('3.5  AI 引擎架构', level=2)
    doc.add_paragraph('边缘推理（场站 II 区 GPU）+ 云端训练（中心云 GPU 集群）+ 人机协同闭环。')

    doc.add_heading('3.6  数据治理体系', level=2)
    doc.add_paragraph('四层治理：主数据管理 → 元数据目录 → 数据血缘 → 数据质量。')

    doc.add_heading('3.7  统一 API 与开发者生态', level=2)
    doc.add_paragraph('北向标准 API (/api/v1/objects/{id}/*) + gRPC（同步）+ Kafka（异步）+ Python/Matlab/C++ SDK。')

    doc.add_heading('3.8  国产化适配', level=2)
    add_word_table(doc,
        ['层次', '推荐', '国产备选', '优先级'],
        [
            ['时序库', 'ClickHouse', 'TDengine/DolphinDB', 'P0'],
            ['业务库', 'Postgres+PostGIS', '达梦 DM8/人大金仓', 'P0'],
            ['OS', 'Linux', '麒麟/统信', 'P1'],
            ['三维引擎', 'Cesium+Three.js（开源）', '—', '✅'],
            ['AI 推理', 'PyTorch/ONNX', '昇思/PaddlePaddle', 'P2'],
        ],
        '表 3-3  国产化适配清单',
        is_color,
    )

    # ──────────────────────────────────────────
    # Chapter 4
    # ──────────────────────────────────────────
    add_chapter(doc, '第四章  业务子系统方案')

    for sub, desc in [
        ('4.1  一体化数智场站平台 (F1)', '全要素数据建模 / 多类型数据处理 / 低代码组态 / 三维可视化 / 高频数据接入'),
        ('4.2  功率-能量协同优化 (F2)', '稳态多源协同 / 暂态快速调频 / 安全防误'),
        ('4.3  构网型透明场站 (F3)', '机组并网性能评价 / 场站并网性能评价 / 发电性能评价'),
        ('4.4  智能运维系统 (F4)', '升压站巡视 / 光伏运维 / 输电线路 / 储能安全 / 运维任务管理'),
        ('4.5  数字运维管理平台 (F5)', '360° 实时监控 / 智能 AI 诊断 / 智慧运维门户 / 生产运营管理'),
    ]:
        doc.add_heading(sub, level=2)
        doc.add_paragraph(desc + '……')

    # ──────────────────────────────────────────
    # Chapter 5
    # ──────────────────────────────────────────
    add_chapter(doc, '第五章  系统非功能特性')

    doc.add_heading('5.1  性能指标承诺', level=2)
    doc.add_paragraph('34 项指标逐项承诺值 + 分层 SLA。量化价值：运维人效 ↑40%+ / 暂态人工介入 ↓80% / 发电损失 ↓15-20%。')

    doc.add_heading('5.2  可靠性设计', level=2)
    doc.add_paragraph('MTBF ≥10000h 的架构保障：N+1 冗余 / 故障自动切换 / 数据多副本。')

    doc.add_heading('5.3  安全性', level=2)
    doc.add_paragraph('等保二级：身份认证 / 访问控制 / 安全审计 / 数据加密 / 入侵检测。电力二次安防：安全 I/II/III 区隔离。')

    doc.add_heading('5.4  可扩展性设计', level=2)
    doc.add_paragraph('多场站统一调度（station_id 隔离）/ 功能模块可插拔 / 低代码建模。')

    doc.add_heading('5.5  开放 API 与开发者生态', level=2)
    doc.add_paragraph('北向 API + gRPC + Kafka + Python/Matlab/C++ 多语言 SDK。')

    # ──────────────────────────────────────────
    # Chapter 6
    # ──────────────────────────────────────────
    add_chapter(doc, '第六章  项目实施计划')

    doc.add_heading('6.1  分阶段交付', level=2)
    add_word_table(doc,
        ['阶段', '周期', '交付内容'],
        [
            ['一阶段', '第 1-2 月', '平台底座（F1）+ 实时监控（F5.1）——跑通数据全链路'],
            ['二阶段', '第 2-3 月', '核心业务（F2/F3/F4 核心模块）——业务功能上线'],
            ['三阶段', '第 3-4 月', '增强功能（P2 项）+ 联调测试 + 上线验收'],
        ],
        '表 6-1  分阶段交付计划',
        is_color,
    )

    doc.add_heading('6.2  里程碑与交付物', level=2)
    add_word_table(doc,
        ['里程碑', '时间', '交付物'],
        [
            ['M1', '月末 1', '平台底座可用，实时数据链路打通'],
            ['M2', '月末 2', '核心业务模块上线，三维可视化可用'],
            ['M3', '月末 3', '全功能联调完成，UAT 测试'],
            ['M4', '月末 4', '正式上线 + 验收文档交付'],
        ],
        '表 6-2  里程碑计划',
        is_color,
    )

    doc.add_heading('6.3  团队配置', level=2)
    add_word_table(doc,
        ['角色', '人数', '职责'],
        [
            ['架构师', '1', '总体架构设计、技术决策'],
            ['前端（三维）', '2', '三维可视化、Web 门户'],
            ['后端（数据平台）', '3', 'USDF 底座、数据处理'],
            ['AI 工程师', '2', 'CV 模型训练与部署'],
            ['电力协议工程师', '1', 'IEC104/GOOSE 集成'],
            ['测试工程师', '1', '全链路压测、验收测试'],
            ['项目经理', '1', '进度管理、客户沟通'],
        ],
        '表 6-3  团队配置',
        is_color,
    )

    # ──────────────────────────────────────────
    # Chapter 7
    # ──────────────────────────────────────────
    add_chapter(doc, '第七章  风险管控')

    doc.add_heading('7.1  技术风险', level=2)
    add_word_table(doc,
        ['#', '风险', '等级', '应对'],
        [
            ['RT1', '百万测点写入瓶颈', '🔴', 'Kafka ≥32 分区 + ClickHouse'],
            ['RT2', '500 并发定义模糊', '🟡', '三层并发模型 + 合同附件'],
            ['RT3', '暂态通信 100 节点', '🟡', '网络专项设计'],
            ['RT4', 'AI 误检率 30%', '🟡', '优化路线图 30%→15%'],
            ['RT5', '端到端 ≤1s 含渲染', '🟡', '分层 SLA 承诺'],
        ],
        '表 7-1  技术风险',
        is_color,
    )

    doc.add_heading('7.2  业务风险', level=2)
    add_word_table(doc,
        ['#', '风险', '等级', '应对'],
        [
            ['RB1', '并发/实时性定义模糊致验收争议', '🔴', '合同附件明确定义'],
            ['RB2', '仿真模块遗漏致报价漏项', '🔴', 'Q2 确认范围'],
            ['RB3', 'AI 指标不达标致验收失败', '🟡', '分阶段达标路线图'],
        ],
        '表 7-2  业务风险',
        is_color,
    )

    doc.add_heading('7.3  实施风险', level=2)
    add_word_table(doc,
        ['#', '风险', '等级', '应对'],
        [
            ['RI1', '仿真模块团队能力缺口', '🔴', '外包/合作'],
            ['RI2', 'GOOSE 协议集成复杂度', '🟡', '电力协议工程师专项'],
            ['RI3', '百万测点接入工期压力', '🟡', '分批接入'],
        ],
        '表 7-3  实施风险',
        is_color,
    )

    # ──────────────────────────────────────────
    # Chapter 8
    # ──────────────────────────────────────────
    add_chapter(doc, '第八章  报价与服务')

    doc.add_heading('8.1  项目报价', level=2)
    add_word_table(doc,
        ['类别', '内容', '备注'],
        [
            ['软件平台', 'USDF 底座 + 五业务域许可', '—'],
            ['定制开发', '三维可视化 / AI 模型 / 接口适配', '按工作量估算'],
            ['三维建模', '风机/光伏/储能/升压站', '按设备数量'],
            ['GIS 数据服务', '地形 / 影像 / 空间分析', '—'],
            ['硬件设备', '服务器 / GPU / 网络 / 安全', '代采购'],
            ['实施与培训', '部署 / 联调 / 用户培训', '—'],
            ['运维质保', '质保期服务', '建议 2 年'],
        ],
        '表 8-1  项目报价结构',
        is_color,
    )

    doc.add_heading('8.2  服务保障', level=2)
    add_word_table(doc,
        ['指标', '承诺'],
        [
            ['质保期', '建议 2 年'],
            ['SLA', 'MTBF ≥10000h（质保期内持续统计）'],
            ['响应时间', '远程 ≤2h / 现场 ≤24h'],
            ['交付物', '需求规格说明书 / 设计文档 / 测试报告 / 用户手册 / 竣工报告'],
        ],
        '表 8-2  服务保障承诺',
        is_color,
    )

    # ──────────────────────────────────────────
    # Revision History (at the END, before appendices)
    # ──────────────────────────────────────────
    add_chapter(doc, '修订记录')

    add_word_table(doc,
        ['版本', '修订日期', '修订内容', '修订人'],
        [
            ['V1.0', '2026-XX-XX', '初稿编制', ''],
            ['', '', '', ''],
            ['', '', '', ''],
        ],
        '表 R-1  修订记录',
        is_color,
    )

    # ── Appendices ──
    add_chapter(doc, '附录')
    doc.add_paragraph('附录 A  术语表')
    doc.add_paragraph('附录 B  参考标准与规范')
    doc.add_paragraph('附录 C  接口规范清单')
    doc.add_paragraph('附录 D  功能完整清单（109 项）')
    doc.add_paragraph('附录 E  非功能指标完整清单（34 项）')
    doc.add_paragraph('附录 F  公司资质与典型案例')


def add_content_b(doc, is_color=False):
    """Content for Template B: 5 technical chapters, no revision history."""
    add_chapter(doc, '第一章  项目概述', is_first=True)

    doc.add_heading('1.1  项目背景', level=2)
    doc.add_paragraph('冀北风光储基地概况、新型电力系统要求、数字化向数智化演进……')

    doc.add_heading('1.2  需求理解', level=2)
    doc.add_paragraph('S (现状) → C (三大挑战) → Q (核心问题) → A (方案概要)……')

    add_image_placeholder(doc, '图 1-1  系统总体架构', 14, 5)

    add_chapter(doc, '第二章  平台建设思想与顶层设计')

    doc.add_heading('2.1  核心思想：空间对象操作系统', level=2)
    doc.add_paragraph('USDF——新能源场站的空间对象操作系统。')

    add_word_table(doc,
        ['OS 概念', 'USDF 对应', '工程含义'],
        [
            ['文件系统', '空间对象注册', '全局唯一 ID + 空间坐标'],
            ['设备驱动', '统一采集适配器', '多协议→统一格式'],
            ['系统调用', '北向 API', '标准 RESTful 接口'],
            ['进程调度', '三通道分流', '按优先级调度'],
        ],
        '表 2-1  操作系统类比',
        is_color,
    )

    doc.add_heading('2.2  四大架构约束', level=2)
    doc.add_paragraph('可靠 / 扩展 / AI 友好 / 实施效能——四个硬约束决定全部技术选型。')

    doc.add_heading('2.3  三条顶层原理', level=2)
    doc.add_paragraph('三通道分流 / 不可越层 / 空间统一维度。')

    doc.add_heading('2.4  方法论：从指标反推架构', level=2)
    doc.add_paragraph('34 项非功能指标 → 架构决策链。')

    add_chapter(doc, '第三章  总体技术架构')

    doc.add_heading('3.1  逻辑架构（四层）', level=2)
    add_image_placeholder(doc, '图 3-1  逻辑架构图', 14, 6)
    doc.add_paragraph('展示层 → 业务层 → 平台层 → 采集层。')

    doc.add_heading('3.2  部署拓扑', level=2)
    doc.add_paragraph('中心云 → 场站边缘（安全 I/II/III 区）→ 设备层。')

    doc.add_heading('3.3  数据分层与流向', level=2)
    doc.add_paragraph('采集 → Kafka → 存储/缓存 → API 的分层流向。')

    doc.add_heading('3.4  关键技术选型', level=2)
    add_word_table(doc,
        ['决策域', '推荐', '备选'],
        [
            ['时序数据库', 'ClickHouse', 'TDengine/DolphinDB'],
            ['业务数据库', 'Postgres+PostGIS', '达梦 DM8'],
            ['消息队列', 'Kafka', 'RocketMQ'],
            ['三维引擎', 'Cesium+Three.js', '待定'],
        ],
        '表 3-1  关键技术选型',
        is_color,
    )

    doc.add_heading('3.5  AI 引擎架构', level=2)
    doc.add_paragraph('边缘推理（场站 II 区 GPU）+ 云端训练（中心云 GPU 集群）+ 人机协同闭环。')

    add_chapter(doc, '第四章  业务子系统方案（概要）')

    for sub, desc in [
        ('4.1  一体化数智场站平台 (F1)', '全要素数据建模 / 三维可视化 / 高频数据接入'),
        ('4.2  功率-能量协同优化 (F2)', '稳态多源协同 / 暂态快速调频 / 安全防误'),
        ('4.3  构网型透明场站 (F3)', '并网性能评价 / 发电性能评价'),
        ('4.4  智能运维系统 (F4)', '升压站巡视 / 光伏运维 / 储能安全'),
        ('4.5  数字运维管理平台 (F5)', '实时监控 / AI 诊断 / 运维门户'),
    ]:
        doc.add_heading(sub, level=2)
        doc.add_paragraph(desc + '……')

    add_chapter(doc, '第五章  系统非功能特性与实施')

    doc.add_heading('5.1  性能指标承诺', level=2)
    doc.add_paragraph('34 项指标逐项承诺值 + 分层 SLA。')

    doc.add_heading('5.2  可靠性设计', level=2)
    doc.add_paragraph('MTBF ≥10000h / N+1 冗余 / 故障自动切换。')

    doc.add_heading('5.3  安全性', level=2)
    doc.add_paragraph('等保二级 + 电力二次安防（I/II/III 区隔离）。')

    doc.add_heading('5.4  可扩展性设计', level=2)
    doc.add_paragraph('多场站统一调度 / 功能模块可插拔 / 低代码建模。')

    doc.add_heading('5.5  项目实施计划', level=2)
    add_word_table(doc,
        ['阶段', '周期', '交付内容'],
        [
            ['一阶段', '第 1-2 月', '平台底座（F1）+ 实时监控'],
            ['二阶段', '第 2-3 月', '核心业务（F2/F3/F4 核心模块）'],
            ['三阶段', '第 3-4 月', '增强功能 + 联调测试 + 上线验收'],
        ],
        '表 5-1  项目实施计划',
        is_color,
    )


# ══════════════════════════════════════════════
#  Template generators
# ══════════════════════════════════════════════

def _setup_doc():
    """Create and configure base document."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.5)
    configure_styles(doc)
    return doc


def generate_template_a(output_dir):
    """Template A: Comprehensive Bid Version."""
    doc = _setup_doc()

    # ── Cover (section 0, no header/footer) ──
    add_cover_page_a(doc)

    # ── Front matter: TOC (section 1, roman numerals) ──
    # Page break then section break
    doc.add_page_break()
    toc_section = add_body_section(doc)
    toc_section.header.is_linked_to_previous = False
    toc_section.footer.is_linked_to_previous = False
    doc.add_heading('目录', level=1)
    doc.add_paragraph('[ 此处插入自动目录：引用 → 目录 → 自动目录 ]')

    # ── Body (section 2, arabic numerals, header/footer) ──
    body_section = add_body_section(doc)
    add_header_footer(body_section, '止善科技  |  冀北风光储数智化生产支撑与数字孪生解决方案')

    add_content_a(doc, is_color=True)

    path = os.path.join(output_dir, 'JBFG-应标版样张.docx')
    doc.save(path)
    print(f"[OK] Template A saved: {path}")


def generate_template_b(output_dir):
    """Template B: Technical Solution Version (gray-friendly)."""
    doc = _setup_doc()

    add_cover_page_b(doc)

    doc.add_page_break()
    toc_section = add_body_section(doc)
    toc_section.header.is_linked_to_previous = False
    toc_section.footer.is_linked_to_previous = False
    doc.add_heading('目录', level=1)
    doc.add_paragraph('[ 此处插入自动目录：引用 → 目录 → 自动目录 ]')

    body_section = add_body_section(doc)
    add_header_footer(body_section, '冀北风光储数智化生产支撑与数字孪生  技术方案')

    add_content_b(doc, is_color=False)

    path = os.path.join(output_dir, 'JBFG-技术方案版样张.docx')
    doc.save(path)
    print(f"[OK] Template B saved: {path}")


if __name__ == '__main__':
    output_dir = r'C:\Users\haoxc\HXC_DATA\80_Knowledges\ZSSolutions\03-项目课题\26.0618-数智化生产支撑系统(JBFG)\20-解决方案\模板体系'
    generate_template_a(output_dir)
    generate_template_b(output_dir)
    print(f"\n[DIR] Output directory: {output_dir}")
    print("[TIP] These are .docx sample files. To create .dotx templates,")
    print("    open in Word and save as 'Word Template (*.dotx)'.")
